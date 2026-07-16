import io


def export_answer_docx(question: str, answer: str, citations: list[dict]) -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Document Q&A Export", level=0)
    document.add_heading("Question", level=1)
    document.add_paragraph(question)
    document.add_heading("Answer", level=1)
    document.add_paragraph(answer)
    if citations:
        document.add_heading("Citations", level=1)
        for i, c in enumerate(citations, 1):
            loc = []
            if c.get("page_number"):
                loc.append(f"page {c['page_number']}")
            if c.get("heading"):
                loc.append(c["heading"])
            loc_str = f" ({', '.join(loc)})" if loc else ""
            document.add_paragraph(
                f"{i}. {c.get('document_name', 'document')}{loc_str} "
                f"[confidence {c.get('confidence', 0)}]",
                style="List Number",
            )
            document.add_paragraph(f"\"{c.get('quoted_text', '')}\"")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def export_comparison_xlsx(comparison: dict) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Comparison"
    docs = comparison["documents"]
    header = ["Topic"] + [d["name"] for d in docs]
    ws.append(header)
    for row in comparison["rows"]:
        line = [row["topic"]]
        for d in docs:
            line.append(row["values"].get(str(d["id"]), ""))
        ws.append(line)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_summary_pdf(title: str, summary: str, citations: list[dict]) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for para in summary.split("\n"):
        if para.strip():
            flow.append(Paragraph(para.replace("\n", "<br/>"), styles["BodyText"]))
            flow.append(Spacer(1, 6))
    if citations:
        flow.append(Spacer(1, 12))
        flow.append(Paragraph("Citations", styles["Heading2"]))
        for i, c in enumerate(citations, 1):
            flow.append(
                Paragraph(
                    f"{i}. {c.get('document_name', 'document')} "
                    f"(p. {c.get('page_number', '-')}) — "
                    f"\"{c.get('quoted_text', '')}\"",
                    styles["BodyText"],
                )
            )
    doc.build(flow)
    return buf.getvalue()
